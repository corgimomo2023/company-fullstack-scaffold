from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / 'docs/reports/ai-vibe-coding-model-comparison-test-report-v4.docx'
OUT.parent.mkdir(parents=True, exist_ok=True)

# DESIGN.md production palette
GREEN = '006A63'
GREEN_DARK = '003531'
GREEN_ACTIVE = '001C19'
ORANGE = 'E6762D'
ORANGE_DARK = 'B15315'
TEXT = '333333'
MUTED = '6C757D'
WHITE = 'FFFFFF'
SURFACE = 'FFFFFF'
SUBTLE = 'F7F7F7'
MUTED_SURFACE = 'ECECEC'
BORDER = 'CECECE'
WARM = 'FFF2EA'
DANGER = 'DC3545'
PALE_GREEN = 'E7F3F1'
PALE_ORANGE = 'FFF2EA'

MODELS = [
    'deepseek-ai/DeepSeek-V4-Flash',
    'deepseek-ai/DeepSeek-V4-Pro',
    'Qwen/Qwen3.5-9B',
]


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn('w:shd'))
    if node is None:
        node = OxmlElement('w:shd')
        tc_pr.append(node)
    node.set(qn('w:val'), 'clear')
    node.set(qn('w:fill'), fill)


def cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn('w:tcMar'))
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        el = tc_mar.find(qn(f'w:{side}'))
        if el is None:
            el = OxmlElement(f'w:{side}')
            tc_mar.append(el)
        el.set(qn('w:w'), str(value))
        el.set(qn('w:type'), 'dxa')


def set_cell(cell, value, *, bold=False, color=TEXT, size=8.2, fill=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(str(value))
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    if fill:
        shade(cell, fill)
    cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(headers, rows, widths=None, header_fill=GREEN_DARK, font_size=8.2):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_repeat_table_header(t.rows[0])
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True, color=WHITE, size=font_size, fill=header_fill)
        if widths:
            t.rows[0].cells[i].width = Cm(widths[i])
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            fill = SUBTLE if ri % 2 else WHITE
            set_cell(cells[i], value, bold=(i == 0), color=TEXT, size=font_size, fill=fill)
            if widths:
                cells[i].width = Cm(widths[i])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    return t


def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(text, *, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def bullets(items, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.03
        p.add_run(item)


def numbered(items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def callout(title, text, fill=PALE_GREEN, title_color=GREEN_DARK):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade(cell, fill)
    cell_margins(cell, 140, 170, 140, 170)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + '\n')
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(title_color)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break():
    doc.add_page_break()


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('Page ')
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.addnext(fld)


def set_section_header_footer(section, label):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    hp = section.header.paragraphs[0]
    hp.text = label
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if hp.runs:
        hp.runs[0].font.name = 'Arial'
        hp.runs[0].font.size = Pt(8)
        hp.runs[0].font.bold = True
        hp.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    fp = section.footer.paragraphs[0]
    fp.text = ''
    add_page_number(fp)


def landscape_section(label):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Cm(1.2)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin = Cm(1.2)
    sec.right_margin = Cm(1.2)
    set_section_header_footer(sec, label)
    return sec


def portrait_section(label):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)
    set_section_header_footer(sec, label)
    return sec


doc = Document()
settings = doc.settings._element
zoom = settings.find(qn('w:zoom'))
if zoom is None:
    zoom = OxmlElement('w:zoom')
    settings.insert(0, zoom)
zoom.set(qn('w:val'), 'bestFit')
zoom.set(qn('w:percent'), '100')
sec = doc.sections[0]
sec.top_margin = Cm(1.6)
sec.bottom_margin = Cm(1.6)
sec.left_margin = Cm(1.7)
sec.right_margin = Cm(1.7)
set_section_header_footer(sec, 'AI VIBE CODING · MODEL COMPARISON · VERSION 4')

styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(9.3)
styles['Normal'].font.color.rgb = RGBColor.from_string(TEXT)
styles['Title'].font.name = 'Arial'
styles['Title'].font.size = Pt(30)
styles['Title'].font.bold = True
styles['Title'].font.color.rgb = RGBColor.from_string(GREEN_DARK)
for name, size, color in [
    ('Heading 1', 18, GREEN_DARK),
    ('Heading 2', 13, GREEN),
    ('Heading 3', 10.5, ORANGE_DARK),
]:
    styles[name].font.name = 'Arial'
    styles[name].font.size = Pt(size)
    styles[name].font.bold = True
    styles[name].font.color.rgb = RGBColor.from_string(color)
if 'Small Note' not in styles:
    st = styles.add_style('Small Note', WD_STYLE_TYPE.PARAGRAPH)
else:
    st = styles['Small Note']
st.font.name = 'Arial'
st.font.size = Pt(8)
st.font.color.rgb = RGBColor.from_string(MUTED)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(64)
r = p.add_run('AI VIBE CODING')
r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(ORANGE_DARK)
p = doc.add_paragraph(style='Title')
p.paragraph_format.space_after = Pt(4)
p.add_run('Model Comparison\nEvaluation Report')
p = doc.add_paragraph()
r = p.add_run('Version 4 · Research-complete methodology · Measured test results intentionally unfilled')
r.font.name = 'Arial'; r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(MUTED)
doc.add_paragraph()
add_table(['Document field', 'V4 value'], [
    ('Decision question', 'Which candidate delivers the most usable accepted application with the least time, cost, token use and human correction?'),
    ('Candidate endpoints', 'DeepSeek-V4-Flash; DeepSeek-V4-Pro; Qwen3.5-9B (exact provider IDs in Section 2)'),
    ('Application A', 'Responsive To-Do List web application'),
    ('Application B', 'Original 1942-style vertical shooter'),
    ('Evaluation layers', 'External evidence 15% · Controlled application test 55% · Human evaluation 30%'),
    ('Status', 'Research and execution design complete; controlled and human results not yet measured'),
    ('Research / pricing date', '2026-08-17'),
    ('Prepared for', 'Internal AI model evaluation team'),
], [4.3, 12.4], font_size=8.8)
callout('V4 interpretation', 'This is not a sparse blank template. The research basis, exact candidate IDs, provider limits, acceptance criteria, test procedure, token definitions, formulas, limitations and decision rules are complete. Only values that require real execution or human observation remain “To be measured”.')
page_break()

# 1 Executive summary
heading('1. Executive Summary')
body('V4 is ready for a controlled model evaluation, but it does not name a winner because no controlled Application A, Application B or human-test measurements have been supplied. Selecting a model now would convert vendor claims into an unsupported internal result.')
callout('Current V4 result', 'Proceed to the controlled test. Do not approve a production model yet. The strongest research finding is that provider endpoint limits differ materially from the underlying open-weight model cards, so the exact endpoint—not the family name—must be the unit of comparison.', PALE_ORANGE, ORANGE_DARK)
add_table(['Management question', 'Research-backed answer before testing', 'Measured result'], [
    ('Which model is best overall?', 'Cannot be decided from model cards; apply the 15/55/30 evidence model and mandatory gates.', 'Not yet measured'),
    ('Which is cheapest?', 'Provider list prices are known, but cost per accepted application depends on real tokens, retries and successful builds.', 'Not yet measured'),
    ('Which is fastest?', 'Parameter size or “Flash” naming does not prove end-to-end completion time in the coding harness.', 'Not yet measured'),
    ('Which supports screenshots?', 'Qwen3.5-9B weights are multimodal, but the current provider endpoint reports image input unsupported; the two DeepSeek endpoints also report no image input.', 'Endpoint canary required'),
    ('Which has the longest usable context?', 'The DeepSeek model cards state 1M and Qwen3.5-9B states 262,144 native; current provider endpoints expose 131,072 and 32,768 respectively.', 'Freeze endpoint limits'),
], [4.0, 8.8, 4.0], font_size=8.5)
heading('1.1 Decision gates', 2)
bullets([
    'All mandatory acceptance criteria must pass in the final accepted build; a high weighted score cannot compensate for a critical functional failure.',
    'No unresolved Severity 1 or Severity 2 defect, secret exposure, unlicensed asset, or critical automated security finding.',
    'Exact provider/model ID, endpoint limits, prompt version, repository commit and usage response must be captured for every run.',
    'Missing measurement is “Not yet measured”, never zero. A model with incomplete evidence is not rankable.',
    'The recommendation is valid only for the frozen versions, harness, prompts, applications and test date.',
])

# 2 Candidate research
page_break()
heading('2. Candidate Model Research')
body('The candidate set is based on the exact models currently exposed by the authenticated provider catalogue used by the existing internal integration. The catalogue response was queried on 2026-08-17 and only non-secret metadata was retained. Prices are USD per one million tokens and are provider list prices, not measured test cost.')
add_table(['Candidate endpoint', 'Official model-card facts', 'Provider endpoint facts', 'Practical implication'], [
    ('deepseek-ai/DeepSeek-V4-Flash', 'MIT; text generation; MoE 284B total / 13B active; model card states 1M context.', 'Max tokens 131,072; input $0.14/M; output $0.28/M; cache read $0.05/M; cache write $0.14/M; image input not supported.', 'Likely cost-focused candidate, but “Flash” does not establish app success or speed. Test text-only baseline.'),
    ('deepseek-ai/DeepSeek-V4-Pro', 'MIT; text generation; MoE 1.6T total / 49B active; model card states 1M context.', 'Max tokens 131,072; input $0.435/M; output $0.87/M; cache read $0.003625/M; cache write $0.435/M; image input not supported.', 'Higher list price and larger active model; test whether quality gain offsets time and cost.'),
    ('Qwen/Qwen3.5-9B', 'Apache-2.0; 9B; causal language model with vision encoder; 262,144 native and extensible to 1,010,000 tokens.', 'Max tokens 32,768; input $0.30/M; output $0.30/M; cache read/write $0.30/M; image input reported unsupported.', 'The weights are multimodal, but this endpoint is currently text-only and shorter-context. Do not claim screenshot support without a canary.'),
], [4.4, 5.2, 5.6, 4.6], font_size=7.6)
heading('2.1 Qwen3.5-Flash clarification', 2)
body('The official Qwen3.5-35B-A3B model card states that hosted Qwen3.5-Flash corresponds to Qwen3.5-35B-A3B and adds production features such as default 1M context and built-in tools. Qwen3.5-35B-A3B is Apache-2.0, has 35B total parameters with about 3B activated, includes a vision encoder and has a native 262,144-token context. It is not currently listed by the tested Pioneer provider catalogue, so it must not be silently substituted for Qwen3.5-9B in the controlled comparison.')
callout('Separate-study rule', 'If the team wants to test Qwen3.5-35B-A3B locally, create a separate self-hosted track with its own GPU, serving, energy and operations cost. Do not mix local inference time or cost with managed API figures as though the environments were equivalent.', PALE_ORANGE, ORANGE_DARK)
heading('2.2 Provider capability verification before Run 1', 2)
add_table(['Check', 'Required action', 'Expected evidence'], [
    ('Model identity', 'Call the provider model catalogue and save a redacted response.', 'Exact ID, provider, endpoint limit, list prices and capability flags'),
    ('Text generation', 'Run a minimal deterministic text canary through the same coding harness.', 'Request/response ID and usage object'),
    ('Image input', 'Run one tiny screenshot canary only where the endpoint claims support.', 'Accepted response or explicit unsupported error'),
    ('Token fields', 'Inspect raw usage JSON before building aggregate formulas.', 'Provider-native field names and inclusion rules'),
    ('Pricing', 'Capture the authenticated catalogue or pricing page on the test date.', 'Price date and currency'),
], [3.2, 8.6, 5.0], font_size=8.3)

# 3 Research basis
page_break()
heading('3. Research Basis and Evidence Hierarchy')
add_table(['Source', 'What it contributes', 'What it cannot prove'], [
    ('ISO/IEC 25010:2023', 'A nine-characteristic product-quality model for requirements and evaluation.', 'It does not prescribe these model weights or declare an AI-model winner.'),
    ('W3C WCAG 2.2', 'Testable web accessibility criteria; V4 targets an automated scan plus selected AA manual checks.', 'Automated tools cannot establish full WCAG conformance.'),
    ('Google Core Web Vitals', 'LCP, INP and CLS definitions and recommended field thresholds.', 'A local Lighthouse run is synthetic and not field-user evidence.'),
    ('OWASP ASVS 5.0.0', 'A basis for testing web application technical security controls.', 'This small evaluation is not a security certification or penetration test.'),
    ('Official model cards', 'Architecture, licence, context and publisher-reported capability.', 'Publisher benchmarks do not prove success in the company coding harness.'),
    ('Authenticated provider catalogue', 'Deployable IDs, endpoint limits, capability flags and list prices on the capture date.', 'It does not prove latency, reliability or accepted-build cost.'),
    ('Controlled internal test', 'Direct evidence for the two frozen applications.', 'Two applications do not represent every enterprise software workload.'),
    ('Human review', 'Usability, clarity, trust and maintainability observations.', 'Subjective scores need common anchors and independent review.'),
], [4.0, 7.0, 6.0], font_size=8.2)
heading('3.1 How public benchmarks may be used', 2)
bullets([
    'Record benchmark name, exact model, reasoning mode, harness, date, sample size and source. Family-level results are not transferable to a different checkpoint or endpoint.',
    'Use coding benchmarks as background only. Algorithmic coding and repository issue repair are narrower than an end-to-end To-Do List or interactive browser game.',
    'Do not merge vendor-reported and independent benchmark values into one average without matching tasks and harnesses.',
    'Do not place external benchmark scores in the 55% controlled-test layer.',
])

# 4 Evaluation model
page_break()
heading('4. Evaluation Architecture')
add_table(['Layer', 'Weight', 'Purpose', 'Evidence required'], [
    ('External evidence', '15%', 'Capability fit, deployability, published limits, pricing and source confidence.', 'Model cards, provider catalogue, dated references'),
    ('Controlled application test', '55%', 'Functional delivery, reliability, autonomy, time, tokens and accepted-build cost.', 'Fresh repositories, logs, automated tests, screenshots, usage JSON'),
    ('Human evaluation', '30%', 'Usability, visual clarity, requirement fit, trust, code readability and correction experience.', 'Common rubric, interview answers, reviewer notes'),
], [4.0, 2.0, 6.4, 5.4], font_size=8.4)
body('These weights are an internal decision rule, not an ISO requirement. They must be frozen before testing. Changing weights after seeing results is prohibited unless the report preserves both the original and revised decision.')
heading('4.1 Controlled-layer subweights', 2)
add_table(['Controlled component', 'Share of controlled layer', 'Measurement'], [
    ('Functional correctness', '45%', 'Mandatory and optional acceptance tests; automated tests'),
    ('Reliability and autonomy', '20%', 'Accepted-run rate, regressions, correction prompts and manual intervention'),
    ('Accepted-build speed', '15%', 'Median time to first runnable and final accepted build'),
    ('Token and cost efficiency', '15%', 'Comparable tokens and cost per accepted build'),
    ('Runtime quality', '5%', 'Web performance / accessibility for App A; FPS / stability for App B'),
], [5.0, 3.4, 8.4], font_size=8.5)
heading('4.2 Human-layer dimensions', 2)
add_table(['Dimension', 'Anchor for 1', 'Anchor for 3', 'Anchor for 5'], [
    ('Task usability', 'Cannot complete core journey', 'Completes with friction', 'Clear and efficient'),
    ('Visual clarity', 'Confusing or broken', 'Usable but inconsistent', 'Clear hierarchy and states'),
    ('Requirement fit', 'Major expectations missing', 'Most needs met', 'All needs met without surprise'),
    ('Trust and recovery', 'Errors are opaque/destructive', 'Recoverable with effort', 'Clear prevention and recovery'),
    ('Maintainability', 'Unsafe or tangled', 'Understandable with issues', 'Well-structured and documented'),
], [4.2, 4.2, 4.2, 4.2], font_size=8.1)

# 5 Fair protocol
page_break()
heading('5. Fair Controlled-Test Protocol')
add_table(['Control', 'Frozen V4 rule', 'Actual setting'], [
    ('Candidate endpoints', 'Use the exact three provider IDs in Section 2; no family aliases.', 'To be confirmed before Run 1'),
    ('Starting repository', 'Equivalent clean repository and dependency lock for each model/app/run.', 'To be recorded'),
    ('Coding harness', 'Same agent/IDE version, system prompt, tools, permissions, retry policy and context policy.', 'To be recorded'),
    ('Environment', 'Same OS, browser, Node/runtime, viewport, hardware and network policy.', 'To be recorded'),
    ('Prompts', 'One frozen prompt per application; identical acceptance criteria and assets.', 'To be versioned'),
    ('Runs', 'Three independent runs per model per app. Report median and range; retain each run.', '18 planned runs'),
    ('Time limit', 'Maximum 60 minutes per run unless frozen otherwise.', '60 minutes proposed'),
    ('Correction limit', 'Maximum two correction prompts after independent defect triage.', 'Two proposed'),
    ('Manual edits', 'Zero before autonomy scoring. Assisted rescue, if any, is a separate non-comparable track.', 'Not permitted in baseline'),
    ('Evidence', 'Same automated test, screenshot, video and log standard for every candidate.', 'Required'),
], [3.6, 9.5, 4.2], font_size=8.2)
heading('5.1 Run sequence', 2)
numbered([
    'Create a fresh workspace from the frozen input commit and assign a unique run ID.',
    'Record model/provider metadata, prompt version, environment and start timestamp.',
    'Submit the frozen application prompt and allow the same tool permissions.',
    'Record first-runnable time; run the identical automated and manual acceptance suite.',
    'An independent reviewer writes a defect list without model-specific hints.',
    'Submit up to two frozen-format correction prompts and rerun the regression suite after each.',
    'Stop at acceptance, time limit or correction limit; record final status and evidence references.',
    'Export provider-native usage fields and pricing evidence before aggregation.',
    'Run the human evaluation on anonymised outputs where practical.',
])
heading('5.2 Failure classification', 2)
add_table(['Class', 'Examples', 'Treatment'], [
    ('Model-quality failure', 'Wrong code, broken command, missed requirement, regression.', 'Counts against model outcome'),
    ('Provider failure', '429, 5xx, service timeout or malformed usage response.', 'Report separately; rerun only under frozen retry rule'),
    ('Harness/infrastructure failure', 'Disk, registry, browser or test-runner outage.', 'Invalidate run if outside model control; preserve evidence'),
    ('Tester intervention', 'Extra hint, manual edit, selective restart.', 'Record and exclude from autonomous baseline'),
], [4.0, 7.0, 6.0], font_size=8.4)

# 6 Scrum plan
page_break()
heading('6. Two-Developer, Five-Day Scrum Execution Plan')
body('Applications are assigned to developers, not models. Each developer runs every candidate on the owned application and cross-reviews the other application. This avoids giving one model a more experienced tester.')
add_table(['Day', 'Developer 1 — App A owner', 'Developer 2 — App B owner', 'Shared exit condition'], [
    ('Day 1 · Freeze', 'Freeze To-Do prompt, tests and starting repo.', 'Freeze shooter prompt, tests, assets and starting repo.', 'Models, environments, prompts, limits and evidence paths signed off'),
    ('Day 2 · Blind runs', 'Run all candidates on App A without corrections.', 'Run all candidates on App B without corrections.', 'First-run logs and automated results captured'),
    ('Day 3 · Repeat + correct', 'Complete remaining runs and correction rounds.', 'Complete remaining runs and correction rounds.', 'Three runs per model/app or documented invalid-run decision'),
    ('Day 4 · Human review', 'Review anonymised App B; collect interview evidence.', 'Review anonymised App A; collect interview evidence.', 'Rubrics complete; disagreements recorded'),
    ('Day 5 · Decision', 'Validate App A evidence and aggregates.', 'Validate App B evidence and aggregates.', 'Independent review, limitation review and recommendation'),
], [2.6, 5.0, 5.0, 5.2], font_size=7.8)
add_table(['Role', 'Accountability'], [
    ('Developer 1', 'Own App A execution and App B cross-review; does not tune only one model.'),
    ('Developer 2', 'Own App B execution and App A cross-review; does not tune only one model.'),
    ('Independent reviewer', 'Verify evidence, severity, rubric anchors and exclusion decisions.'),
    ('Business owner', 'Approve frozen weights and accept or reject the recommendation.'),
], [4.2, 12.8], font_size=8.6)

# 7 App A
page_break()
heading('7. Application A — To-Do List')
body('Purpose: measure common business UI delivery, state management, validation, persistence, responsive behaviour, accessibility, tests and maintainable component structure.')
heading('7.1 Frozen requirement', 2)
body('Build a responsive To-Do List web application. A user can create, edit, delete, complete and reopen tasks; filter all, active and completed tasks; validate empty or whitespace-only input; and preserve tasks after reload. Include an empty state, destructive-action confirmation, keyboard operation, mobile and desktop layouts, automated tests and clear run instructions. Use no backend unless the frozen prompt explicitly requires one.')
app_a = [
    ('A01 Start-up', 'Installs and starts using documented commands without code repair.', 'Mandatory'),
    ('A02 Create task', 'A valid task is created and shown immediately.', 'Mandatory'),
    ('A03 Validation', 'Empty or whitespace-only task is rejected with a clear message.', 'Mandatory'),
    ('A04 Edit task', 'Task text and permitted fields can be edited and saved.', 'Mandatory'),
    ('A05 Complete / reopen', 'State changes correctly and remains consistent.', 'Mandatory'),
    ('A06 Delete', 'Confirmation is shown and confirmed deletion persists.', 'Mandatory'),
    ('A07 Filtering', 'All, active and completed filters return correct results.', 'Mandatory'),
    ('A08 Persistence', 'Tasks remain after reload and reopening.', 'Mandatory'),
    ('A09 Empty state', 'Useful empty-state guidance appears with no tasks.', 'Optional'),
    ('A10 Responsive', 'No clipped controls or horizontal overflow at 390×844 and 1440×900.', 'Mandatory'),
    ('A11 Keyboard', 'All actions are keyboard reachable with visible focus.', 'Mandatory'),
    ('A12 Accessibility', 'No critical automated issue; labels, names and contrast are reviewed.', 'Mandatory'),
    ('A13 Error recovery', 'Corrupt local data fails safely with a recovery path.', 'Optional'),
    ('A14 Automated tests', 'Core create/edit/delete/filter/persistence tests pass.', 'Mandatory'),
    ('A15 Console', 'No uncaught console error in tested journeys.', 'Mandatory'),
    ('A16 Maintainability', 'Reasonable component separation; no secrets or avoidable duplication.', 'Mandatory'),
    ('A17 Performance', 'Record Lighthouse and available LCP/INP/CLS observations.', 'Measured, not a hard field-data claim'),
]
add_table(['ID', 'Expected outcome frozen before testing', 'Gate', 'Actual result', 'Evidence'],
          [(a, b, c, 'Not yet tested', 'To be captured') for a, b, c in app_a],
          [3.0, 7.6, 3.1, 3.2, 3.2], font_size=7.2)
heading('7.2 App A measurement notes', 2)
bullets([
    'Synthetic Lighthouse data is a repeatable lab measurement, not a substitute for real-user field data.',
    'Automated accessibility tools find only a subset of issues; keyboard, focus order, labels and visible states need manual checks.',
    'Local-storage behaviour must be tested across reload and corrupt-data recovery, not inferred from source code.',
])

# 8 App B
page_break()
heading('8. Application B — Original 1942-Style Vertical Shooter')
body('Purpose: measure interactive state, game-loop correctness, collision handling, animation, keyboard/touch input, original asset use, runtime performance, stability and maintainable separation of systems.')
heading('8.1 Frozen requirement', 2)
body('Build an original browser-based vertical shooter inspired by 1940s aircraft arcade pacing. Do not copy protected game art, names, levels, music or sounds. Include one stage, one boss, continuous movement, player fire, at least three enemy behaviours, collisions and damage, score, lives or health, at least two power-ups, pause, restart, game-over, mute, keyboard and touch controls, responsive play, automated tests and clear run instructions.')
app_b = [
    ('B01 Start-up', 'Opens to a playable start screen without code repair.', 'Mandatory'),
    ('B02 Movement', 'Keyboard and touch movement are responsive and bounded.', 'Mandatory'),
    ('B03 Firing', 'Player fire works with sensible rate limiting.', 'Mandatory'),
    ('B04 Enemy variety', 'At least three visibly different enemy behaviours work.', 'Mandatory'),
    ('B05 Collision', 'Player, enemy and projectile collisions register once and correctly.', 'Mandatory'),
    ('B06 Damage / lives', 'Damage, invulnerability and death state are consistent.', 'Mandatory'),
    ('B07 Score', 'Score changes by defined events and resets on restart.', 'Mandatory'),
    ('B08 Power-ups', 'At least two bounded power-up effects work.', 'Mandatory'),
    ('B09 Stage and boss', 'Stage progression reaches an original boss encounter.', 'Mandatory'),
    ('B10 Pause / resume', 'Gameplay and timers pause and resume without corruption.', 'Mandatory'),
    ('B11 Game over / restart', 'Restart creates a clean state without duplicate loops.', 'Mandatory'),
    ('B12 Audio / mute', 'Original/licensed audio or placeholders are documented; mute works.', 'Optional'),
    ('B13 Responsive play', 'Playable at 390×844 and 1440×900 without clipped HUD.', 'Mandatory'),
    ('B14 Performance', 'Record median FPS, 1% low and long/dropped frames on fixed hardware.', 'Mandatory'),
    ('B15 Stability', 'Ten-minute play session has no crash or runaway resource growth.', 'Mandatory'),
    ('B16 Originality', 'No copied protected art, audio, branding, maps or levels.', 'Mandatory'),
    ('B17 Automated tests', 'Scoring, collision, reset and bounds tests pass.', 'Mandatory'),
    ('B18 Maintainability', 'Loop, entities, input, audio, UI and configuration are separated.', 'Mandatory'),
]
add_table(['ID', 'Expected outcome frozen before testing', 'Gate', 'Actual result', 'Evidence'],
          [(a, b, c, 'Not yet tested', 'To be captured') for a, b, c in app_b],
          [3.0, 7.6, 3.1, 3.2, 3.2], font_size=7.2)
heading('8.2 App B measurement notes', 2)
bullets([
    'Record FPS on the same device, browser, viewport and power profile. Do not compare one model on desktop with another on mobile.',
    'A brief smooth-looking demo is not enough; run the same scripted sequence and ten-minute stability session.',
    'Original mechanics are permitted, but protected game assets, names, maps, soundtracks and sprites are out of scope.',
])

# 9 Expected vs actual
page_break()
heading('9. Expected-versus-Actual Result Method')
body('Expected outcomes are completed in advance in Sections 7 and 8. “Actual result” must describe observed behaviour, not restate the expectation. Every failed or partial result must cite evidence and severity.')
add_table(['Field', 'Completion rule', 'Example format without inventing a result'], [
    ('Expected outcome', 'Frozen before model runs and linked to a requirement ID.', '“Task remains after browser reload.”'),
    ('Actual status', 'Pass / Partial / Fail / Not executed / Invalid run.', 'Not yet tested'),
    ('Observed behaviour', 'One factual sentence describing what happened.', 'To be observed'),
    ('Evidence', 'Automated report, screenshot, video, console log or commit reference.', 'To be captured'),
    ('Severity', 'S1 critical, S2 major, S3 moderate, S4 minor.', 'To be assigned if defective'),
    ('Correction outcome', 'Fixed / partially fixed / not fixed / regression introduced.', 'Not yet tested'),
], [4.0, 8.0, 5.0], font_size=8.4)
heading('9.1 Accepted-build definition', 2)
body('An accepted build passes every mandatory acceptance item, has no unresolved S1 or S2 defect, remains inside the frozen time and correction limits, has no manual code repair in the baseline, and has complete evidence and usage capture. A high partial-test percentage is not an accepted build.')
heading('9.2 Invalid-run rule', 2)
body('A run may be invalidated only for a documented provider or infrastructure failure outside model control. Preserve the original run and exclusion reason. Do not discard a poor result because it is inconvenient, and do not rerun one model under a more favourable prompt.')

# 10 landscape run capture
landscape_section('AI VIBE CODING · V4 RUN AND TOKEN CAPTURE')
heading('10. Run-Level Measurement Ledger')
body('Complete every row after the final correction turn. Keep provider-native usage JSON as evidence. The tables below cover three candidates × two applications × three runs.')
run_rows = []
for model in MODELS:
    for app in ['A — To-Do List', 'B — Vertical Shooter']:
        for run in [1, 2, 3]:
            run_rows.append((model, app, str(run), 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'Not yet tested', 'To be captured'))
add_table(['Exact model ID', 'Application', 'Run', 'First runnable', 'Accepted time', 'Correction prompts', 'Manual interventions', 'Final status', 'Evidence ref'],
          run_rows, [5.2, 3.8, 1.2, 2.6, 2.6, 2.8, 3.2, 2.7, 3.2], font_size=6.6)
heading('10.1 Token, cache, image and cost ledger', 2)
usage_rows = []
for model in MODELS:
    for app in ['A', 'B']:
        for run in [1, 2, 3]:
            usage_rows.append((model, app, str(run), 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'To be captured'))
add_table(['Exact model ID', 'App', 'Run', 'Input tokens', 'Output tokens', 'Reasoning tokens', 'Cache read', 'Cache write', 'Image count / units', 'Provider total', 'API cost', 'Usage evidence'],
          usage_rows, [4.8, 1.0, 1.0, 2.2, 2.2, 2.3, 2.1, 2.1, 2.5, 2.2, 1.9, 3.1], font_size=6.3)
heading('10.2 Additional execution measurements', 2)
add_table(['Measurement', 'Definition', 'Entry rule'], [
    ('Run ID / timestamp', 'Unique ID and UTC start/end time.', 'Required for every run'),
    ('Prompt version', 'Hash or version of frozen task and correction template.', 'Required'),
    ('Provider request IDs', 'IDs for initial and correction requests where exposed.', 'Preserve in evidence'),
    ('Automated tests', 'Pass/fail/skip count and exact command.', 'Do not write only “tests passed”'),
    ('First-runnable time', 'Submission to first successful local start.', 'Include failed start attempts'),
    ('Accepted-build time', 'Submission to accepted state including corrections.', 'Undefined when no accepted build'),
    ('Provider errors', '429/5xx/timeout count and charged status.', 'Separate from model failure'),
    ('Evidence reference', 'Restricted path/URI to logs, usage JSON, screenshots and video.', 'No credentials or customer data'),
], [4.0, 11.0, 10.0], font_size=7.5)

# 11 token dictionary portrait
portrait_section('AI VIBE CODING · V4 TOKEN ACCOUNTING')
heading('11. Token and Cost Data Dictionary')
add_table(['Field', 'Definition', 'Normalisation rule'], [
    ('Input / prompt tokens', 'Provider-reported tokens for prompts, context, tool results and correction turns.', 'Sum initial and correction requests; preserve raw fields.'),
    ('Output / completion tokens', 'Provider-reported generated content.', 'Do not assume reasoning is separate.'),
    ('Reasoning / thinking tokens', 'Separately exposed internal reasoning tokens, if any.', 'If included in output, do not add again.'),
    ('Cached input / read', 'Tokens served from provider cache.', 'Keep separate and apply cache-read price.'),
    ('Cache write', 'Tokens charged for cache population where exposed.', 'Keep separate and apply cache-write price.'),
    ('Image count / image tokens', 'Images, pixels, tiles or token-equivalent units reported by provider.', 'Record provider unit; do not invent image tokens.'),
    ('Provider total tokens', 'Provider’s reported total.', 'Preserve unchanged even when it differs from a simple sum.'),
    ('Comparable total', 'A documented sum of mutually exclusive token fields.', 'Calculate only after confirming inclusion rules.'),
    ('API cost', 'Request-level cost or price-calculated marginal API cost.', 'Record pricing date and currency.'),
    ('Tool/runtime cost', 'Sandbox, browser, storage or hosted-agent charges outside model usage.', 'Report separately before combined cost.'),
    ('Effective subscription cost', 'Allocated share of a fixed monthly plan.', 'Not directly comparable with API token billing without method.'),
], [4.1, 7.4, 5.3], font_size=8.0)
callout('Do not double-count', 'Providers expose usage differently. Some include reasoning in output; some expose cache reads and writes separately; some do not expose image tokens; hosted-agent tool calls may be outside standard token fields. Never compare one unexplained “total tokens” number without preserving raw provider fields and pricing date.', PALE_ORANGE, ORANGE_DARK)
heading('11.1 Cost formulas', 2)
bullets([
    'Marginal model cost = input tokens × input rate + output tokens × output rate + cache/image/reasoning charges according to provider rules.',
    'Accepted-build cost = total valid-run cost ÷ accepted builds. If accepted builds = 0, report undefined/infinite, not zero.',
    'Tokens per accepted build = comparable mutually exclusive tokens ÷ accepted builds.',
    'Do not treat unused subscription capacity as zero cost. Present API marginal cost and subscription allocation in separate columns.',
])

# 12 aggregate tables
landscape_section('AI VIBE CODING · V4 APPLICATION AGGREGATES')
heading('12. Application-Level and Combined Result Tables')
summary_rows = [(m, 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'To be measured') for m in MODELS]
heading('12.1 Application A — To-Do List totals', 2)
add_table(['Exact model ID', 'Valid runs', 'Accepted builds', 'Input', 'Output', 'Comparable total', 'Total cost', 'Tokens / accepted', 'Cost / accepted'],
          summary_rows, [5.5, 2.2, 2.5, 2.3, 2.3, 3.0, 2.5, 3.2, 3.0], font_size=7.2)
heading('12.2 Application B — Vertical Shooter totals', 2)
add_table(['Exact model ID', 'Valid runs', 'Accepted builds', 'Input', 'Output', 'Comparable total', 'Total cost', 'Tokens / accepted', 'Cost / accepted'],
          summary_rows, [5.5, 2.2, 2.5, 2.3, 2.3, 3.0, 2.5, 3.2, 3.0], font_size=7.2)
heading('12.3 Combined comparison', 2)
add_table(['Exact model ID', 'A accepted', 'B accepted', 'Combined comparable tokens', 'Combined cost', 'Median accepted time', 'Corrections / accepted', 'Mandatory gate', 'Rank'],
          [(m, 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'To be measured', 'Not yet evaluated', 'Not rankable') for m in MODELS],
          [5.2, 2.3, 2.3, 3.7, 2.5, 3.3, 3.4, 3.1, 2.2], font_size=7.0)
heading('12.4 Why application totals remain separate', 2)
body('App A and App B exercise different failure modes and token patterns. A combined total alone can hide a model that succeeds on the simple business UI but fails on real-time game logic. Always publish App A, App B and combined views together.')

# 13 human eval
portrait_section('AI VIBE CODING · V4 HUMAN EVALUATION')
heading('13. Human Evaluation')
body('Human testers evaluate anonymised builds where practical. Every reviewer uses the same task script and 1–5 anchored rubric. Record the rating, the factual observation supporting it and any reviewer disagreement.')
add_table(['Question / task', 'Evidence sought', 'Response'], [
    ('Can you create and manage a task without instruction?', 'First-use clarity and task completion friction.', 'Not yet collected'),
    ('Which To-Do control or label was confusing?', 'Specific navigation, wording or state problem.', 'Not yet collected'),
    ('Did the To-Do result match the written requirements?', 'Expectation mismatch and missing features.', 'Not yet collected'),
    ('Could you recover from invalid input or deletion?', 'Error prevention, explanation and recovery.', 'Not yet collected'),
    ('Could you play the shooter on keyboard and touch?', 'Control discoverability, responsiveness and mobile usability.', 'Not yet collected'),
    ('Were game status, damage, score and boss progress clear?', 'HUD clarity and state feedback.', 'Not yet collected'),
    ('Did pause, game over and restart behave predictably?', 'Trust in state transitions and recovery.', 'Not yet collected'),
    ('Which build would you release, and why?', 'Release confidence tied to observable evidence.', 'Not yet collected'),
    ('Which codebase would you maintain next month?', 'Structure, naming, tests and documentation.', 'Not yet collected'),
    ('What one improvement is most important?', 'Actionable priority, not generic preference.', 'Not yet collected'),
], [6.0, 7.2, 3.8], font_size=8.0)
heading('13.1 Reviewer scoring sheet', 2)
add_table(['Reviewer', 'Model code', 'App', 'Usability', 'Visual clarity', 'Requirement fit', 'Trust/recovery', 'Maintainability', 'Observation ref'], [
    ('To be assigned', 'Blind code', 'A/B', 'Not scored', 'Not scored', 'Not scored', 'Not scored', 'Not scored', 'To be captured'),
    ('To be assigned', 'Blind code', 'A/B', 'Not scored', 'Not scored', 'Not scored', 'Not scored', 'Not scored', 'To be captured'),
], [2.6, 2.3, 1.0, 1.7, 2.0, 2.0, 2.0, 2.1, 3.0], font_size=7.4)
body('User interview answers are evidence within the end-to-end case study; they are not a substitute for executing the complete application workflow and acceptance suite.')

# 14 scoring
page_break()
heading('14. Scoring and Decision Rules')
body('Scores are calculated only after the evidence set is complete. Keep unrounded component values and round only displayed totals. Missing data does not become zero and cannot be ranked.')
add_table(['Layer / rule', 'Formula or decision'], [
    ('External evidence score', '0–100 from capability fit, deployability, verified limits/pricing and source confidence.'),
    ('Controlled test score', '0–100 from the frozen subweights in Section 4.1.'),
    ('Human evaluation score', '0–100 from anchored reviewer ratings and resolved evidence.'),
    ('Overall score', '0.15 × External + 0.55 × Controlled + 0.30 × Human.'),
    ('Mandatory gate', 'Any final mandatory failure, unresolved S1/S2, secret exposure or unlicensed asset makes the candidate Not Acceptable.'),
    ('Tie rule', 'Prefer higher accepted-build rate, then lower human assistance, then lower cost per accepted build.'),
    ('Instability rule', 'Report all runs, median and range; do not hide failed runs behind an average.'),
], [5.0, 12.0], font_size=8.5)
heading('14.1 Scorecard', 2)
add_table(['Exact model ID', 'External 15%', 'Controlled 55%', 'Human 30%', 'Mandatory gate', 'Overall', 'Decision'], [
    (m, 'Not yet scored', 'Not yet measured', 'Not yet measured', 'Not yet evaluated', 'Not rankable', 'Pending evidence') for m in MODELS
], [4.8, 2.3, 2.7, 2.5, 2.8, 2.2, 2.8], font_size=7.5)
callout('No premature winner', 'V4 deliberately does not reuse simulated V3 values or earlier unrelated benchmark runs. Only measurements produced under this frozen V4 protocol may populate the result tables.')

# 15 limitations
page_break()
heading('15. Limitations and Decision Impact')
add_table(['Limitation', 'Practical impact', 'Required response'], [
    ('Only two applications', 'Results may not transfer to backend, mobile, data pipelines, enterprise integration or large repositories.', 'Run a second study for materially different workloads.'),
    ('Three runs per model/app', 'Exploratory medians and ranges do not establish a stable population difference.', 'Use at least ten runs when differences drive procurement.'),
    ('Provider endpoints differ from checkpoints', 'Context, vision, tools and pricing may be service-layer constrained.', 'Freeze endpoint metadata and avoid family-level claims.'),
    ('Developer effect', 'Prompt handling and defect triage may influence outcomes.', 'Assign apps, cross-review outputs and use fixed intervention rules.'),
    ('Provider load and network', 'Latency and errors can change by time and region.', 'Timestamp runs and report provider errors separately.'),
    ('Synthetic performance tests', 'Lighthouse and scripted FPS do not represent all users/devices.', 'Keep claims local to the fixed environment.'),
    ('Automated accessibility/security', 'Tools detect only a subset of issues.', 'Perform manual checks; do not claim certification.'),
    ('Model and price drift', 'A later endpoint may not reproduce the result.', 'Record date, version, raw metadata and pricing evidence.'),
    ('Self-hosted vs API cost', 'Hardware, energy and operations costs differ from token prices.', 'Use a separate cost model and study track.'),
    ('Human scoring', 'Reviewer preference and familiarity introduce subjectivity.', 'Use anchored rubrics, multiple reviewers and observations.'),
], [4.2, 7.3, 5.5], font_size=8.0)
heading('15.1 Claims that V4 does not support', 2)
bullets([
    '“Model X is generally the best coding model.”',
    '“A 1M model card means the current endpoint accepts 1M tokens.”',
    '“The cheapest token price produces the cheapest accepted application.”',
    '“No automated warning means the application is accessible or secure.”',
    '“A visually polished first screen is a completed case study.”',
])

# 16 case study
page_break()
heading('16. End-to-End Case Study Structure')
body('The V4 case study is the complete decision journey, not a set of isolated Q&A examples. It is completed after the controlled runs using the sections below and links to evidence rather than copying large logs into the report.')
add_table(['Case-study stage', 'Completed V4 content', 'Execution result field'], [
    ('1. Business question', 'Select a practical model for two representative vibe-coding workflows.', 'Confirmed'),
    ('2. Candidate research', 'Exact model IDs, model-card facts, endpoint limits, capability gaps and prices documented.', 'Completed 2026-08-17'),
    ('3. Frozen requirements', 'App A and App B expected outcomes and mandatory gates defined.', 'Completed'),
    ('4. Fair execution', 'Same repositories, prompts, harness, limits, evidence and correction policy.', 'To be executed'),
    ('5. Defect correction', 'Independent defect list and maximum two common-format correction prompts.', 'To be executed'),
    ('6. Automated validation', 'Functional, regression, accessibility, performance and stability evidence.', 'To be executed'),
    ('7. Human evaluation', 'Anonymised workflows, anchored rubric, interview answers and observations.', 'To be executed'),
    ('8. Token/cost analysis', 'Run-level provider usage, App A totals, App B totals and per-accepted-build efficiency.', 'To be measured'),
    ('9. Limitations', 'Decision impact and non-generalisation boundaries stated before recommendation.', 'Completed; update after runs'),
    ('10. Decision', 'Mandatory gate first, then weighted score and tie rules.', 'Pending complete evidence'),
], [4.3, 9.3, 3.5], font_size=8.0)
heading('16.1 Final case-study narrative template', 2)
body('“The team evaluated [exact model IDs] because [business need]. Under the frozen environment and prompts, each candidate attempted the To-Do List and original vertical shooter three times. [MODEL] produced [measured outcome] with [measured correction effort], while [MODEL] produced [measured outcome]. Human reviewers observed [evidence-backed observation]. After applying the mandatory gates and 15/55/30 scoring rule, [MODEL] was [recommended / not recommended] for [bounded use]. The decision does not extend beyond the tested endpoint versions, applications, tools and date.”')

# 17 evidence/signoff
page_break()
heading('17. Evidence Pack, Review and Sign-off')
add_table(['Evidence item', 'Required content', 'Status'], [
    ('Frozen prompts', 'Version/hash for App A, App B and correction prompt.', 'To be captured'),
    ('Source state', 'Input commit and separate output commit per run.', 'To be captured'),
    ('Environment', 'OS, browser, runtime, hardware, harness and permissions.', 'To be captured'),
    ('Model metadata', 'Exact provider/model ID, endpoint limit, capability flags and price date.', 'Research captured; reconfirm at Run 1'),
    ('Usage data', 'Raw provider usage JSON and derived formulas.', 'To be captured'),
    ('Automated tests', 'Commands, pass/fail/skip counts and reports.', 'To be captured'),
    ('Visual evidence', 'Fixed viewport screenshots and shooter recording.', 'To be captured'),
    ('Human evidence', 'Rubric scores, answers, observations and disagreement.', 'To be captured'),
    ('Decision calculation', 'Layer scores, gates, medians, ranges and exclusions.', 'To be calculated'),
], [4.0, 9.5, 3.5], font_size=8.2)
add_table(['Approval role', 'Name', 'Date', 'Decision / comment'], [
    ('Developer 1 / App A owner', 'To be assigned', 'To be entered', 'Pending'),
    ('Developer 2 / App B owner', 'To be assigned', 'To be entered', 'Pending'),
    ('Independent reviewer', 'To be assigned', 'To be entered', 'Pending'),
    ('Business owner', 'To be assigned', 'To be entered', 'Pending'),
], [4.6, 4.2, 3.4, 4.8], font_size=8.5)

# 18 sources
page_break()
heading('18. Sources and Research Notes')
sources = [
    ('[1] DeepSeek official Hugging Face model card — DeepSeek-V4-Flash', 'https://huggingface.co/deepseek-ai/DeepSeek-V4-Flash', 'Architecture, licence, context and publisher evaluation context; accessed 2026-08-17.'),
    ('[2] DeepSeek official Hugging Face model card — DeepSeek-V4-Pro', 'https://huggingface.co/deepseek-ai/DeepSeek-V4-Pro', 'Architecture, licence, context and publisher evaluation context; accessed 2026-08-17.'),
    ('[3] Qwen official Hugging Face model card — Qwen3.5-9B', 'https://huggingface.co/Qwen/Qwen3.5-9B', 'Vision encoder, 9B model, native/extended context and licence; accessed 2026-08-17.'),
    ('[4] Qwen official Hugging Face model card — Qwen3.5-35B-A3B', 'https://huggingface.co/Qwen/Qwen3.5-35B-A3B', 'Qwen3.5-Flash correspondence, 35B/3B MoE, vision and context; accessed 2026-08-17.'),
    ('[5] Pioneer authenticated model catalogue', 'https://api.pioneer.ai/v1/models', 'Exact deployable IDs, endpoint limits, capability flags and list prices; redacted authenticated response captured 2026-08-17.'),
    ('[6] ISO/IEC 25010:2023', 'https://www.iso.org/standard/78176.html', 'Product quality model and lifecycle evaluation use.'),
    ('[7] W3C Web Content Accessibility Guidelines 2.2', 'https://www.w3.org/TR/WCAG22/', 'Current web accessibility recommendation and testable criteria.'),
    ('[8] Google web.dev — Web Vitals', 'https://web.dev/articles/vitals', 'User-centred loading, responsiveness and visual-stability metrics.'),
    ('[9] OWASP Application Security Verification Standard', 'https://owasp.org/www-project-application-security-verification-standard/', 'ASVS 5.0.0 basis for web technical security controls.'),
]
for title, url, note in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title + '\n')
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN_DARK)
    p.add_run(url + '\n')
    r = p.add_run(note)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    r.font.size = Pt(8.5)
heading('18.1 Research confidence labels', 2)
add_table(['Label', 'Meaning'], [
    ('Verified external fact', 'Directly stated by an official model card, standard or live provider catalogue.'),
    ('Internal rule', 'A frozen evaluation decision such as weights, run count or correction limit.'),
    ('To be measured', 'Requires controlled execution, provider usage output or human observation.'),
    ('Not comparable', 'Different model, endpoint, harness, environment or billing basis.'),
], [4.4, 12.6], font_size=8.6)

# Core properties
props = doc.core_properties
props.title = 'AI Vibe Coding Model Comparison Evaluation Report — Version 4'
props.subject = 'Research-complete evaluation methodology for DeepSeek V4 Flash, DeepSeek V4 Pro and Qwen3.5-9B'
props.author = 'Internal AI Model Evaluation Team'
props.keywords = 'vibe coding, model comparison, DeepSeek V4, Qwen3.5, To-Do List, 1942 shooter, token usage, human evaluation'
props.comments = 'No controlled application-test or human-test figures are fabricated. Only published/provider facts and frozen methodology are populated.'

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
