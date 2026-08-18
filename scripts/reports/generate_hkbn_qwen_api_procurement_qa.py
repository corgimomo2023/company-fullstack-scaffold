from pathlib import Path
from datetime import date
from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs/reports/hkbn-qwen-model-api-procurement-vendor-qa.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

GREEN = "006A63"
GREEN_DARK = "003531"
ORANGE = "E6762D"
ORANGE_DARK = "B15315"
TEXT = "333333"
MUTED = "6C757D"
WHITE = "FFFFFF"
SUBTLE = "F7F7F7"
BORDER = "CECECE"
PALE_GREEN = "E7F3F1"
PALE_ORANGE = "FFF2EA"
DANGER = "DC3545"
TODAY = date(2026, 8, 18).isoformat()


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:val"), "clear")
    node.set(qn("w:fill"), fill)


def cell_margins(cell, top=70, start=85, bottom=70, end=85):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tc_mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def set_cell(cell, text, *, bold=False, color=TEXT, fill=None, size=7.8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    if fill:
        shade(cell, fill)
    cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def table(headers, rows, widths, font_size=7.8, first_col_bold=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    repeat_header(t.rows[0])
    for i, header in enumerate(headers):
        set_cell(t.rows[0].cells[i], header, bold=True, color=WHITE, fill=GREEN_DARK, size=font_size)
        set_cell_width(t.rows[0].cells[i], widths[i])
    for row_index, row in enumerate(rows):
        cells = t.add_row().cells
        fill = WHITE if row_index % 2 == 0 else SUBTLE
        for i, value in enumerate(row):
            set_cell(cells[i], value, bold=(first_col_bold and i == 0), fill=fill, size=font_size)
            set_cell_width(cells[i], widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.03
        p.add_run(item)


def callout(title, text, fill=PALE_GREEN, title_color=GREEN_DARK):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    shade(cell, fill)
    cell_margins(cell, 140, 170, 140, 170)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "\n")
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(title_color)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def set_header_footer(section, label):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    hp = section.header.paragraphs[0]
    hp.text = label
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if hp.runs:
        hp.runs[0].font.name = "Arial"
        hp.runs[0].font.size = Pt(8)
        hp.runs[0].font.bold = True
        hp.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    fp = section.footer.paragraphs[0]
    fp.text = ""
    page_number(fp)


def landscape(label):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    sec.top_margin = Cm(1.15)
    sec.bottom_margin = Cm(1.15)
    sec.left_margin = Cm(1.2)
    sec.right_margin = Cm(1.2)
    set_header_footer(sec, label)
    return sec


def supplier_rows(section_code, questions):
    return [
        (f"{section_code}.{i}", question, "Supplier to complete", "Attach evidence or cite contract / SLA clause")
        for i, question in enumerate(questions, 1)
    ]


SECTIONS = [
    ("1", "Supplier, Service and Commercial Scope", [
        "Identify the contracting legal entity, company registration jurisdiction and the party responsible for service delivery and support.",
        "Confirm whether HKBN operates the service directly or resells / integrates another provider. List every material subcontractor and its role.",
        "State the service name, orderable SKU, edition and geographic region proposed for this purchase.",
        "Describe whether the offer is public cloud API, private endpoint, dedicated deployment, on-premises deployment or another model.",
        "State the proposed contract term, minimum commitment, renewal mechanism, termination rights and early-termination charges.",
        "List all one-off charges, implementation fees, support fees, platform fees, minimum monthly spend and overage charges.",
        "Confirm billing currency, exchange-rate mechanism, tax treatment and whether prices are fixed for the contract term.",
        "Provide a full price-validity period and the notice period for any price or packaging change.",
        "Describe available trial, proof-of-concept and benchmark arrangements, including duration, token credit, support and restrictions.",
        "Provide three relevant enterprise customer references or anonymised case studies for managed LLM API use.",
    ]),
    ("2", "Exact Model Identity and Deployment", [
        "Provide the exact API model ID and marketing name. Family-level answers such as 'Qwen' or 'Qwen Max' are not sufficient.",
        "Provide the upstream publisher, model checkpoint / revision, release date and model-card URL.",
        "State whether the endpoint is base, instruct, reasoning, multimodal or a vendor-modified derivative.",
        "Disclose any fine-tuning, alignment, system prompt, safety layer, routing, speculative decoding or model substitution applied by the service.",
        "State total and active parameter counts where applicable, architecture type and supported modalities.",
        "State serving precision and quantisation exactly, including weight, activation and KV-cache precision (for example BF16, FP8, INT8 or INT4).",
        "Identify the serving engine and material runtime version (for example vLLM or another implementation).",
        "Confirm whether inference uses dedicated or shared compute, and identify the accelerator class without exposing sensitive infrastructure details.",
        "Describe multi-tenancy isolation and whether another customer's workload can affect latency or capacity.",
        "Confirm the physical processing region(s), failover region(s), data-storage region(s) and whether any processing occurs outside Hong Kong.",
        "Describe how the endpoint behaves during model upgrades and whether customers can pin a model version.",
        "State model deprecation, end-of-life and migration notice periods, including rollback options after a change.",
    ]),
    ("3", "API, Compatibility and Functional Features", [
        "Provide the production and non-production API base URLs and API versioning policy.",
        "List supported OpenAI-compatible endpoints, including /v1/chat/completions, /v1/responses, model catalogue and embeddings where applicable.",
        "Confirm streaming support and protocol details, including SSE event format, termination event and error behaviour.",
        "Provide authentication methods, key scopes, key rotation, expiry and separate credentials for development, UAT and production.",
        "Confirm support for tool / function calling, parallel tool calls, tool-choice control and the exact request / response schema.",
        "Confirm structured output / JSON schema support and explain validation or retry behaviour when output is invalid.",
        "Confirm image input support, accepted formats, size / count limits and image-token metering. If unsupported, state this explicitly.",
        "Confirm reasoning controls, reasoning-token visibility and whether reasoning tokens are returned, hidden or billable.",
        "Describe system-message handling, prompt precedence, content filtering and any fields that are silently ignored or rewritten.",
        "State maximum request body size, timeout limits, streaming idle timeout and maximum generation duration.",
        "Confirm support for idempotency keys, request IDs, trace IDs and customer-supplied metadata that is not used as prompt content.",
        "Provide error-code documentation, retry guidance and Retry-After behaviour for throttling and transient failures.",
        "Confirm availability of asynchronous batch inference, batch pricing and maximum batch size where applicable.",
        "Provide official SDKs, supported languages, sample code and compatibility limitations against the OpenAI SDK.",
    ]),
    ("4", "Context Window, Tokens and Usage Accounting", [
        "State the upstream model's native context window and the endpoint's enforced context window separately.",
        "State maximum input tokens, maximum output tokens, combined context limit and any lower defaults.",
        "Confirm whether the endpoint can support approximately one million tokens and whether this uses native capability, YaRN scaling or another extension.",
        "If extended context is offered, disclose supported scaling factors, expected quality trade-offs, latency impact and pricing impact.",
        "Identify the tokenizer and provide a supported method or library for pre-flight token counting.",
        "Explain truncation behaviour when input exceeds the limit: reject, truncate oldest, truncate newest or another policy.",
        "Define all usage response fields for input, output, reasoning, cached, image and tool-call tokens.",
        "State whether system prompts, tool schemas, retrieved context and repeated conversation history are billable input tokens.",
        "Explain prompt-cache eligibility, cache key / scope, TTL, isolation, invalidation and minimum cacheable length.",
        "Confirm whether cached reads and writes are visible in every API response and billing export.",
        "Provide a reconciled sample showing raw request, usage JSON and resulting invoice calculation with confidential values removed.",
    ]),
    ("5", "Performance, Capacity and Service Levels", [
        "Provide measured time-to-first-token (TTFT) at P50, P95 and P99 for representative short, medium and long prompts.",
        "Provide time-per-output-token (TPOT) and output tokens per second at P50, P95 and P99.",
        "Provide end-to-end latency at P50, P95 and P99 for at least three declared input/output token profiles.",
        "State the benchmark period, region, concurrency, sample size, streaming mode and whether results include queuing and cold starts.",
        "Provide performance at concurrency 1, 5, 10, 25 and the proposed production concurrency, or explain supported alternatives.",
        "State hard and soft limits for requests per minute, tokens per minute, concurrent requests, daily quota and maximum burst.",
        "Explain whether limits apply per key, user, tenant, IP, model, endpoint or account, and whether input and output tokens share one quota.",
        "Describe capacity reservation, committed throughput and the process / lead time for increasing limits.",
        "State the monthly availability SLA, measurement method, exclusions, maintenance treatment and service-credit schedule.",
        "Provide historical monthly availability and major incident history for the proposed service for the latest available 12 months.",
        "Define planned maintenance windows, advance notice and whether maintenance can reduce capacity without counting as downtime.",
        "Define severity levels, acknowledgement time, update frequency, target restoration time and root-cause-analysis delivery time.",
        "Describe overload behaviour, queue limits, admission control, 429 / 503 handling and protection from noisy neighbours.",
        "Confirm whether a status page, incident subscription and customer-specific service-health feed are available.",
    ]),
    ("6", "Pricing, Metering and Cost Control", [
        "Provide unit price per one million input tokens and output tokens for the exact proposed model ID.",
        "Provide separate prices for reasoning tokens, cached reads, cache writes, image input, tool calls and batch inference.",
        "Confirm whether failed, cancelled, timed-out, filtered or partially streamed requests are charged and how usage is calculated.",
        "Explain rounding rules, minimum billable units and whether token counts are rounded per request or in aggregate.",
        "Provide volume tiers, committed-use discounts, reserved-capacity prices and overage rates.",
        "Provide a monthly worked example for low, expected and peak usage, showing every charge and assumption.",
        "Confirm availability of near-real-time usage dashboards and machine-readable exports by model, key, project and cost centre.",
        "Describe budget alerts, hard spending caps, quota controls and the delay between consumption and reporting.",
        "State invoice dispute procedure, metering audit rights and retention period for detailed usage records.",
        "List support, onboarding, network, private connectivity, storage, logging and professional-service charges not included in token prices.",
        "Confirm the process and notice period for price changes, and whether the customer may terminate without penalty after a material increase.",
    ]),
    ("7", "Data Protection, Security and Compliance", [
        "Confirm contractually that customer prompts, outputs, files and metadata are not used to train or improve any model unless the customer explicitly opts in in writing.",
        "State retention periods separately for prompts, outputs, uploaded files, abuse-monitoring copies, application logs, security logs, backups and billing metadata.",
        "Confirm whether zero-retention or customer-configurable retention is available and identify any exceptions.",
        "Describe deletion, backup expiry and verifiable data-erasure procedures at request and contract termination.",
        "Identify data controller / processor roles and provide the proposed Data Processing Agreement.",
        "List all subprocessors, processing locations, data-transfer mechanisms and notification / objection rights for subprocessor changes.",
        "Confirm whether any prompt, output or metadata crosses into Mainland China or another jurisdiction, including support and telemetry access.",
        "State encryption in transit standards and encryption at rest standards, including key ownership and customer-managed-key options.",
        "Describe tenant isolation, network segmentation, private connectivity, IP allowlisting, VPN / private link and egress controls.",
        "Describe RBAC, SSO / SAML / OIDC, MFA, administrative separation and least-privilege controls for the management portal.",
        "Describe audit-log contents, integrity, export method, retention and integration with customer SIEM.",
        "Describe secret management, API-key storage, rotation, revocation and controls preventing credentials from appearing in logs.",
        "State whether supplier personnel can access customer content, under what conditions, with what approval, logging and confidentiality controls.",
        "Provide current ISO 27001, ISO 27701, SOC 2 Type II, CSA STAR or equivalent certifications and scope statements.",
        "Provide vulnerability-management, penetration-testing, secure-development and dependency / model-supply-chain controls.",
        "Describe prompt-injection, data-exfiltration, abuse, malware and harmful-content protections, including controls that customers can configure.",
        "State security-incident notification time, investigation / evidence commitments, regulatory cooperation and liability provisions.",
        "Describe business continuity, disaster recovery, backup, RTO and RPO for control plane, inference service and customer configuration.",
        "State governing law, regulatory-access handling and the process for responding to government or law-enforcement requests.",
        "Confirm the customer's rights to export configurations, usage logs and relevant data in a usable format before termination.",
    ]),
    ("8", "Support and Operating Model", [
        "Describe included support tier, support hours, languages, channels and Hong Kong public-holiday coverage.",
        "Provide named escalation routes for technical, security, commercial and service-management issues.",
        "State response and restoration targets by severity and whether 24x7 support is available for Severity 1 incidents.",
        "Describe onboarding, solution architecture, performance tuning and migration support included in the proposal.",
        "Confirm release-note, breaking-change and model-behaviour-change notification mechanisms.",
        "Describe change management for API schema, safety policy, model checkpoint, quantisation, runtime and infrastructure changes.",
        "Confirm access to non-production / sandbox environments and whether their limits, model and behaviour match production.",
        "Provide the service review cadence and reports covering availability, latency, incidents, capacity, usage and cost.",
        "Describe exit assistance, data return / deletion, configuration export and migration support at contract end.",
    ]),
]


doc = Document()
settings = doc.settings._element
zoom = settings.find(qn("w:zoom"))
if zoom is None:
    zoom = OxmlElement("w:zoom")
    settings.insert(0, zoom)
zoom.set(qn("w:val"), "bestFit")
zoom.set(qn("w:percent"), "100")

sec = doc.sections[0]
sec.top_margin = Cm(1.6)
sec.bottom_margin = Cm(1.6)
sec.left_margin = Cm(1.7)
sec.right_margin = Cm(1.7)
set_header_footer(sec, "HKBN QWEN MODEL API · PROCUREMENT Q&A · DRAFT")

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(9.3)
styles["Normal"].font.color.rgb = RGBColor.from_string(TEXT)
styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(28)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor.from_string(GREEN_DARK)
for name, size, color in (("Heading 1", 17, GREEN_DARK), ("Heading 2", 12.5, GREEN), ("Heading 3", 10.5, ORANGE_DARK)):
    styles[name].font.name = "Arial"
    styles[name].font.size = Pt(size)
    styles[name].font.bold = True
    styles[name].font.color.rgb = RGBColor.from_string(color)
if "Small Note" not in styles:
    small = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
else:
    small = styles["Small Note"]
small.font.name = "Arial"
small.font.size = Pt(8)
small.font.color.rgb = RGBColor.from_string(MUTED)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
r = p.add_run("ENTERPRISE AI PROCUREMENT")
r.font.name = "Arial"
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(ORANGE_DARK)
p = doc.add_paragraph(style="Title")
p.paragraph_format.space_after = Pt(5)
p.add_run("HKBN Qwen Model API\nVendor Due-Diligence Q&A")
p = doc.add_paragraph()
r = p.add_run("Technical · Commercial · Performance · Security · Trial Acceptance")
r.font.name = "Arial"
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string(MUTED)
doc.add_paragraph()
table(["Document field", "Value"], [
    ("Status", "Draft for supplier response — not an order or production approval"),
    ("Version", "1.0"),
    ("Prepared date", TODAY),
    ("Prepared for", "Internal AI Platform, Procurement, Information Security and Legal review"),
    ("Proposed supplier", "HKBN / contracting entity to be confirmed by supplier"),
    ("Subject", "Managed Qwen model API and associated support services"),
    ("Response deadline", "[TO BE COMPLETED]"),
    ("Commercial validity", "[SUPPLIER TO COMPLETE]"),
], [4.5, 12.2], font_size=8.8)
callout("Decision rule", "No purchase decision should rely on a family name, headline context size or average latency alone. The exact endpoint, checkpoint, serving precision, limits, billing rules, data handling, SLA and measured trial evidence must be contractually identified.")

# Purpose and instructions
doc.add_page_break()
heading("Purpose and Response Instructions")
body("This questionnaire is designed to support technical, commercial, security and legal due diligence for a managed Qwen model API proposed by HKBN. It deliberately separates upstream model capability from the capability and operating limits of the purchased endpoint.")
heading("Supplier response instructions", 2)
bullets([
    "Answer every question against the exact service and model SKU proposed. Do not answer only at Qwen-family level.",
    "Use 'Not supported', 'Not available' or 'Not disclosed' where applicable; do not leave material fields ambiguous.",
    "Attach official documentation, test methodology, certification, sample usage JSON, price sheet and proposed contract / SLA clauses.",
    "Mark roadmap functions separately from generally available contracted functions, with an expected date and dependency.",
    "Identify any answer that differs between trial, UAT and production.",
    "Redact credentials, customer data, internal hostnames and other secrets from all evidence.",
])
heading("Mandatory information before commercial approval", 2)
table(["Gate", "Required evidence", "Status"], [
    ("Exact endpoint identity", "Exact model ID, checkpoint / revision, region, modality, context and version-pinning policy", "Open"),
    ("Transparent serving profile", "Precision / quantisation, dedicated or shared compute, runtime and multi-tenant isolation", "Open"),
    ("Measured service performance", "P50 / P95 / P99 TTFT, TPOT, output TPS and end-to-end latency under declared concurrency", "Open"),
    ("Complete pricing", "Input, output, reasoning, cache, image, batch, minimum spend, support and all non-token charges", "Open"),
    ("Data-protection commitment", "No training without written opt-in; retention, residency, subprocessor and deletion commitments", "Open"),
    ("Contracted service levels", "Availability, support severity targets, incident notice, service credits and change notice", "Open"),
    ("Controlled trial", "Representative workload, raw usage evidence, reliability, cost reconciliation and security canaries", "Open"),
], [4.2, 10.8, 2.0], font_size=8.3)
callout("Important", "All thresholds in the acceptance plan are proposed for joint confirmation before the trial. Missing evidence must be recorded as a procurement risk, not silently treated as a pass.", PALE_ORANGE, ORANGE_DARK)

# Main questionnaire
landscape("HKBN QWEN MODEL API · SUPPLIER QUESTIONNAIRE")
heading("Supplier Questionnaire")
body("Supplier response and evidence columns are intentionally left open for completion. Continue answers in an annex where necessary and reference the question ID.")
for index, (code, title, questions) in enumerate(SECTIONS):
    if index:
        doc.add_page_break()
    heading(f"{code}. {title}")
    table(
        ["ID", "Question", "Supplier response", "Evidence / contract reference"],
        supplier_rows(code, questions),
        [1.25, 11.6, 7.0, 6.6],
        font_size=7.45,
    )

# Proposed trial plan
landscape("HKBN QWEN MODEL API · TRIAL AND ACCEPTANCE")
heading("9. Proposed Trial and Acceptance Plan")
body("The final test profiles and thresholds must be frozen jointly before execution. The trial should use the exact production-intended model ID and region, with raw request IDs and usage records retained in redacted form.")
table(["Area", "Proposed test", "Evidence", "Pass criterion to freeze before trial"], [
    ("Identity canary", "Call the model catalogue and a minimal deterministic prompt.", "Model catalogue, response model field, request ID", "Exact endpoint and checkpoint are identifiable; no silent substitution."),
    ("API compatibility", "Exercise streaming, non-streaming, system prompt, structured output and tool calling.", "Request / response samples and automated contract tests", "Required fields work; unsupported fields are documented; errors are stable."),
    ("Token accounting", "Run pre-counted prompts with output caps, reasoning and cache where offered.", "Tokenizer count, raw usage JSON and bill calculation", "Usage is explainable and reconciles to quoted prices within agreed tolerance."),
    ("Short workload", "Approximately 1k input tokens and 256 output tokens at concurrency 1 and 10.", "TTFT, TPOT, output TPS, E2E P50/P95/P99", "Thresholds agreed after supplier baseline; no hidden queue exclusion."),
    ("Medium workload", "Approximately 16k input tokens and 1k output tokens at representative concurrency.", "Same metrics plus errors and throttling", "Performance and error rate meet agreed business target."),
    ("Long-context workload", "Representative 64k+ input and, only if offered, extended-context test.", "Accuracy checks, latency, memory / limit errors and cost", "No silent truncation; declared limit and quality trade-off are confirmed."),
    ("Load and throttling", "Ramp concurrency and token rate through expected and peak levels.", "429 / 503 rate, Retry-After, recovery and capacity graphs", "Limits match contract; graceful overload; no sustained cross-tenant degradation."),
    ("Reliability", "Repeat frozen task set across multiple days and time bands.", "Success rate, error classes, retry count and latency distribution", "Agreed success and error-rate targets are met, not only average latency."),
    ("Quality", "Use a fixed internal prompt set with objective checks and blinded human review.", "Expected answers, rubric, reviewer notes and failure examples", "Mandatory task gates pass; score is reproducible for the frozen version."),
    ("Security", "Verify key scope / revocation, access controls, redacted logs and prohibited data handling.", "Portal screenshots, audit export and supplier confirmation", "No secret leakage; required controls and contractual data terms are in place."),
    ("Cost", "Reconcile trial usage at low, expected and peak profiles.", "Usage export and shadow invoice", "No unexplained charge category; forecast is within approved budget envelope."),
], [3.3, 8.8, 7.0, 7.35], font_size=7.5)
heading("9.1 Trial controls", 2)
bullets([
    "Use synthetic or approved non-sensitive data until the DPA and security review are complete.",
    "Freeze prompt versions, model ID, parameters, concurrency, endpoint region and test code commit.",
    "Capture timestamps, HTTP status, request ID, usage fields, retry count and client-observed latency for every request.",
    "Report mean only as a secondary statistic; primary latency views are P50, P95 and P99 with sample size.",
    "Do not compare a dedicated trial endpoint with a shared production proposal unless the difference is declared and quantified.",
])

# Schedules
landscape("HKBN QWEN MODEL API · COMMERCIAL AND SLA SCHEDULES")
heading("10. Commercial Schedule Template")
table(["Charge item", "Unit", "Supplier price", "Commitment / tier", "Notes and exclusions"], [
    ("Input tokens", "HKD or USD per 1M tokens", "", "", "Exact model ID required"),
    ("Output tokens", "HKD or USD per 1M tokens", "", "", ""),
    ("Reasoning tokens", "Per 1M tokens / included", "", "", "Visibility and billing rule"),
    ("Prompt-cache read", "Per 1M cached tokens", "", "", "TTL / eligibility"),
    ("Prompt-cache write", "Per 1M tokens", "", "", ""),
    ("Image input", "Per image / token equivalent", "", "", "Resolution / tile rule"),
    ("Batch inference", "Per 1M tokens", "", "", "Turnaround target"),
    ("Reserved capacity", "Monthly / throughput unit", "", "", "Guaranteed capacity"),
    ("Platform / minimum fee", "Monthly", "", "", ""),
    ("Implementation", "One-off", "", "", "Deliverables"),
    ("Support", "Monthly / annual", "", "", "Coverage and severity"),
    ("Private connectivity", "Monthly / one-off", "", "", ""),
    ("Other", "Specify", "", "", "No unlisted charge accepted"),
], [5.0, 4.5, 4.2, 5.6, 7.0], font_size=7.8)
doc.add_page_break()
heading("11. SLA Schedule Template")
table(["Service indicator", "Supplier commitment", "Measurement method", "Exclusions", "Remedy / credit"], [
    ("Monthly availability", "", "", "", ""),
    ("P95 TTFT — agreed profile", "", "", "", ""),
    ("P99 TTFT — agreed profile", "", "", "", ""),
    ("P95 end-to-end latency", "", "", "", ""),
    ("Minimum output throughput", "", "", "", ""),
    ("Committed RPM / TPM / concurrency", "", "", "", ""),
    ("Severity 1 acknowledgement", "", "", "", ""),
    ("Severity 1 restoration target", "", "", "", ""),
    ("Security incident notification", "", "", "", ""),
    ("Model / API breaking-change notice", "", "", "", ""),
    ("Model deprecation notice", "", "", "", ""),
], [5.2, 5.0, 6.2, 5.0, 4.9], font_size=7.8)

# Risk and sign-off
landscape("HKBN QWEN MODEL API · RISK AND SIGN-OFF")
heading("12. Clarification and Deviation Log")
table(["Ref", "Question / requirement", "Supplier deviation", "Risk / impact", "Owner", "Resolution / due date"], [
    ("D-01", "", "", "", "", ""),
    ("D-02", "", "", "", "", ""),
    ("D-03", "", "", "", "", ""),
    ("D-04", "", "", "", "", ""),
    ("D-05", "", "", "", "", ""),
    ("D-06", "", "", "", "", ""),
], [1.4, 6.0, 5.7, 5.7, 3.2, 4.3], font_size=7.8)
heading("13. Internal Review and Sign-off")
table(["Review area", "Reviewer / function", "Outcome", "Conditions / open risks", "Date"], [
    ("Technical architecture and API", "", "Pending", "", ""),
    ("Performance and capacity", "", "Pending", "", ""),
    ("Information security", "", "Pending", "", ""),
    ("Privacy / data protection", "", "Pending", "", ""),
    ("Legal / contract", "", "Pending", "", ""),
    ("Commercial / procurement", "", "Pending", "", ""),
    ("Business owner", "", "Pending", "", ""),
], [5.2, 5.0, 3.0, 9.0, 3.1], font_size=8.0)
callout("Document status", "This questionnaire is a due-diligence working document. A completed response does not itself constitute contract acceptance, security approval or authority to process production data.", PALE_ORANGE, ORANGE_DARK)

# Document metadata hygiene
core = doc.core_properties
core.title = "HKBN Qwen Model API Vendor Due-Diligence Q&A"
core.subject = "Technical, commercial, performance, security and trial questionnaire"
core.author = "Internal AI Platform Team"
core.keywords = "HKBN, Qwen, API, procurement, SLA, security, vendor due diligence"
core.comments = "Draft questionnaire; contains no credentials or production secrets."

doc.save(OUT)
print(OUT)
